# routes/affidavit.py
import uuid
import logging
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from flask import Blueprint, request, jsonify, session, send_file
from docx import Document

from models import db, AffidavitRecord, User
from config import AFFIDAVIT_CONFIG
from routes.auth import admin_required
from helpers.docx_helpers import replace_text_in_paragraph, replace_text_in_tables
from helpers.html_helpers import process_paragraph_html, process_table_html, generate_print_html_page
from services.pdf_service import (
    convert_all_docx_to_pdfs_batch,
    create_print_job,
    cleanup_print_job,
    cancel_print_job
)

logger = logging.getLogger(__name__)
affidavit_bp = Blueprint('affidavit', __name__)

@affidavit_bp.route('/api/affidavits', methods=['GET'])
@admin_required
def get_affidavits():
    try:
        records = AffidavitRecord.query.order_by(AffidavitRecord.modified_at.desc()).all()
        return jsonify({'success': True, 'records': [r.to_dict() for r in records]})
    except Exception as e:
        logger.error(f"Error fetching affidavits: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>', methods=['GET'])
@admin_required
def get_affidavit(id):
    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            return jsonify({'success': False, 'message': 'Record not found'}), 404
        return jsonify({'success': True, 'record': record.to_dict()})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits', methods=['POST'])
@admin_required
def create_affidavit():
    try:
        data = request.get_json()
        template_key = data.get('template_key')
        replacements = data.gfet('replacements', {})

        if not template_key or template_key not in AFFIDAVIT_CONFIG:
            return jsonify({'success': False, 'message': 'Invalid template type'}), 400

        # Ensure deponent/subject name is populated correctly from new name
        if 'NEW_NAME' in replacements and 'UPDATE_NAME' not in replacements:
            replacements['UPDATE_NAME'] = replacements['NEW_NAME']

        # Fallback child name to new name automatically to fit Word template requirements
        if 'NEW_NAME' in replacements and 'CHILD_NAME' not in replacements:
            replacements['CHILD_NAME'] = replacements['NEW_NAME']

        # Smart pronoun derivation for HE_SHE / HIS_HER from SON-DAUGHTER
        sd_val = (replacements.get('SON-DAUGHTER') or '').lower().strip()
        if sd_val == 'son':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif sd_val == 'daughter':
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        # Fallback pronoun derivation for adults based on relationship mappings
        rel_val = (replacements.get('UPDATE_RELATION') or '').lower().strip()
        if rel_val == 's/o':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif rel_val in ['d/o', 'w/o']:
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        primary_name = (
            replacements.get('UPDATE_NAME') or 
            replacements.get('CHILD_NAME') or 
            replacements.get('OLD_NAME') or 
            replacements.get('LANDLORD_NAME') or 
            'UNNAMED'
        ).upper().strip()

                # Smart field processing: auto ALPHA_DATE from NUM_DATE
        if 'NUM_DATE' in replacements and replacements['NUM_DATE']:
            try:
                num_date_val = replacements['NUM_DATE']
                # Handle both YYYY-MM-DD and DD/MM/YYYY formats
                if '-' in num_date_val:
                    dt = datetime.strptime(num_date_val, '%Y-%m-%d')
                elif '/' in num_date_val:
                    dt = datetime.strptime(num_date_val, '%d/%m/%Y')
                else:
                    dt = None

                if dt:
                    day = dt.day
                    suffix = 'TH'
                    if day % 10 == 1 and day % 100 != 11: suffix = 'ST'
                    elif day % 10 == 2 and day % 100 != 12: suffix = 'ND'
                    elif day % 10 == 3 and day % 100 != 13: suffix = 'RD'
                    months = ['JANUARY','FEBRUARY','MARCH','APRIL','MAY','JUNE',
                              'JULY','AUGUST','SEPTEMBER','OCTOBER','NOVEMBER','DECEMBER']
                    replacements['ALPHA_DATE'] = f"{day}{suffix} DAY OF {months[dt.month-1]} {dt.year}"
            except Exception as e:
                logger.warning(f"Failed to auto-generate ALPHA_DATE: {e}")

        # Smart pronoun derivation for HE_SHE / HIS_HER from SON-DAUGHTER
        sd_val = (replacements.get('SON-DAUGHTER') or '').lower().strip()
        if sd_val == 'son':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif sd_val == 'daughter':
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        from helpers.text_helpers import format_date_to_ddmmyyyy
        for field in AFFIDAVIT_CONFIG[template_key]['fields']:
            if field['type'] == 'date' and field['id'] in replacements:
                replacements[field['id']] = format_date_to_ddmmyyyy(replacements[field['id']])

        record = AffidavitRecord(
            created_by=session.get('user_id'),
            template_key=template_key,
            primary_name=primary_name,
            replacements=replacements
        )
        
        db.session.add(record)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Affidavit saved successfully!', 'record': record.to_dict()})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating affidavit: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>', methods=['PUT'])
@admin_required
def update_affidavit(id):
    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            return jsonify({'success': False, 'message': 'Record not found'}), 404

        data = request.get_json()
        template_key = record.template_key
        replacements = data.get('replacements', {})

        # Ensure deponent/subject name is populated correctly from new name
        if 'NEW_NAME' in replacements and 'UPDATE_NAME' not in replacements:
            replacements['UPDATE_NAME'] = replacements['NEW_NAME']

        # Fallback child name to new name automatically to fit Word template requirements
        if 'NEW_NAME' in replacements and 'CHILD_NAME' not in replacements:
            replacements['CHILD_NAME'] = replacements['NEW_NAME']

         # Smart pronoun derivation for HE_SHE / HIS_HER from SON-DAUGHTER
        sd_val = (replacements.get('SON-DAUGHTER') or '').lower().strip()
        if sd_val == 'son':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif sd_val == 'daughter':
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        # Fallback pronoun derivation for adults based on relationship mappings
        rel_val = (replacements.get('UPDATE_RELATION') or '').lower().strip()
        if rel_val == 's/o':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif rel_val in ['d/o', 'w/o']:
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        primary_name = (
            replacements.get('UPDATE_NAME') or 
            replacements.get('CHILD_NAME') or 
            replacements.get('OLD_NAME') or 
            replacements.get('LANDLORD_NAME') or 
            'UNNAMED'
        ).upper().strip()

                # Smart field processing: auto ALPHA_DATE from NUM_DATE
        if 'NUM_DATE' in replacements and replacements['NUM_DATE']:
            try:
                num_date_val = replacements['NUM_DATE']
                # Handle both YYYY-MM-DD and DD/MM/YYYY formats
                if '-' in num_date_val:
                    dt = datetime.strptime(num_date_val, '%Y-%m-%d')
                elif '/' in num_date_val:
                    dt = datetime.strptime(num_date_val, '%d/%m/%Y')
                else:
                    dt = None

                if dt:
                    day = dt.day
                    suffix = 'TH'
                    if day % 10 == 1 and day % 100 != 11: suffix = 'ST'
                    elif day % 10 == 2 and day % 100 != 12: suffix = 'ND'
                    elif day % 10 == 3 and day % 100 != 13: suffix = 'RD'
                    months = ['JANUARY','FEBRUARY','MARCH','APRIL','MAY','JUNE',
                              'JULY','AUGUST','SEPTEMBER','OCTOBER','NOVEMBER','DECEMBER']
                    replacements['ALPHA_DATE'] = f"{day}{suffix} DAY OF {months[dt.month-1]} {dt.year}"
            except Exception as e:
                logger.warning(f"Failed to auto-generate ALPHA_DATE: {e}")

        # Smart pronoun derivation for HE_SHE / HIS_HER from SON-DAUGHTER
        sd_val = (replacements.get('SON-DAUGHTER') or '').lower().strip()
        if sd_val == 'son':
            replacements['HE_SHE'] = 'he'
            replacements['HIS_HER'] = 'his'
        elif sd_val == 'daughter':
            replacements['HE_SHE'] = 'she'
            replacements['HIS_HER'] = 'her'

        from helpers.text_helpers import format_date_to_ddmmyyyy
        for field in AFFIDAVIT_CONFIG[template_key]['fields']:
            if field['type'] == 'date' and field['id'] in replacements:
                replacements[field['id']] = format_date_to_ddmmyyyy(replacements[field['id']])

        record.primary_name = primary_name
        record.replacements = replacements
        record.modified_at = datetime.now(timezone.utc)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Affidavit updated successfully!', 'record': record.to_dict()})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating affidavit: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>', methods=['DELETE'])
@admin_required
def delete_affidavit(id):
    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            return jsonify({'success': False, 'message': 'Record not found'}), 404

        db.session.delete(record)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Affidavit deleted successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>/download', methods=['GET'])
@admin_required
def download_affidavit_docx(id):
    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            return jsonify({'success': False, 'message': 'Record not found'}), 404

        config = AFFIDAVIT_CONFIG.get(record.template_key)
        if not config:
            return jsonify({'success': False, 'message': 'Template config missing'}), 400

        template_path = Path(config['folder']) / config['file']
        if not template_path.exists():
            return jsonify({'success': False, 'message': f'Template file not found at {template_path}'}), 404

        doc = Document(str(template_path))
        replacements = {str(k): (str(v).upper() if v is not None else '') for k, v in (record.replacements or {}).items()}

        for p in doc.paragraphs:
            replace_text_in_paragraph(p, replacements)
        if doc.tables:
            replace_text_in_tables(doc.tables, replacements)

        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_text_in_paragraph(p, replacements)
            if section.header.tables:
                replace_text_in_tables(section.header.tables, replacements)
            for p in section.footer.paragraphs:
                replace_text_in_paragraph(p, replacements)
            if section.footer.tables:
                replace_text_in_tables(section.footer.tables, replacements)

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        filename = f"{config['category'].replace(' ', '_')}_{record.primary_name.replace(' ', '_')}.docx"
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error downloading docx: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>/print-pdf', methods=['GET'])
@admin_required
def print_affidavit_pdf(id):
    job_id = uuid.uuid4().hex
    create_print_job(job_id)

    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            cleanup_print_job(job_id)
            return jsonify({'success': False, 'message': 'Affidavit record not found'}), 404

        config = AFFIDAVIT_CONFIG.get(record.template_key)
        if not config:
            cleanup_print_job(job_id)
            return jsonify({'success': False, 'message': 'Template config missing'}), 400

        template_path = Path(config['folder']) / config['file']
        if not template_path.exists():
            cleanup_print_job(job_id)
            return jsonify({'success': False, 'message': f'Template file not found at {template_path}'}), 404

        stem_name = template_path.stem
        docx_files_dict = {stem_name: template_path}
        replacements = {str(k): (str(v).upper() if v is not None else '') for k, v in (record.replacements or {}).items()}

        try:
            pdf_results = convert_all_docx_to_pdfs_batch(
                docx_files_dict, replacements, job_id=job_id
            )
        except RuntimeError as e:
            cleanup_print_job(job_id)
            if str(e) == 'CANCELLED':
                return jsonify({'success': False, 'cancelled': True, 'message': 'Print was cancelled'}), 200
            raise e

        if stem_name not in pdf_results:
            cleanup_print_job(job_id)
            return jsonify({'success': False, 'message': 'Failed to compile PDF'}), 500

        pdf_bytes = pdf_results[stem_name]
        filename = f"Affidavit_{record.primary_name.replace(' ', '_')}.pdf"

        response = send_file(
            BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=False,
            download_name=filename
        )
        response.headers['X-Print-Job-Id'] = job_id
        return response

    except Exception as e:
        logger.error(f"Error compiling PDF: {str(e)}", exc_info=True)
        cleanup_print_job(job_id)
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cleanup_print_job(job_id)

@affidavit_bp.route('/api/affidavits/<id>/cancel-print', methods=['POST'])
@admin_required
def cancel_affidavit_print(id):
    try:
        data = request.get_json() or {}
        job_id = data.get('job_id')
        if not job_id:
            return jsonify({'success': False, 'message': 'No job_id provided'}), 400
        success, message = cancel_print_job(job_id)
        if not success:
            return jsonify({'success': False, 'message': message}), 404
        return jsonify({'success': True, 'message': message})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@affidavit_bp.route('/api/affidavits/<id>/preview-html', methods=['GET'])
@admin_required
def preview_affidavit_html(id):
    try:
        record = db.session.get(AffidavitRecord, id)
        if not record:
            return "<h3>Record not found</h3>", 404

        config = AFFIDAVIT_CONFIG.get(record.template_key)
        template_path = Path(config['folder']) / config['file']
        if not template_path.exists():
            return "<h3>Template file not found</h3>", 404

        doc = Document(str(template_path))
        replacements = {str(k): (str(v).upper() if v is not None else '') for k, v in (record.replacements or {}).items()}

        all_elements_html = []
        doc_html = ['<div class="print-document-wrapper"><div class="print-document-body">']

        for p in doc.paragraphs:
            para_html = process_paragraph_html(p, replacements)
            if para_html:
                doc_html.append(para_html)

        for t in doc.tables:
            table_html = process_table_html(t, replacements)
            if table_html:
                doc_html.append(table_html)

        doc_html.append('</div></div>')
        all_elements_html.append(''.join(doc_html))

        return generate_print_html_page(all_elements_html, record.primary_name)
    except Exception as e:
        logger.error(f"Error compiling HTML preview: {str(e)}")
        return f"<h3>Error compiling preview: {str(e)}</h3>", 500
